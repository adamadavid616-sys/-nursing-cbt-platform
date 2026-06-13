import hashlib
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path("C:/Users/A_D/Downloads")
QUESTION_OUT = ROOT / "src" / "supplemental-bank.js"
GUIDE_OUT = ROOT / "src" / "guide-content.js"


FILES = [
    "CBT - NUTRITION AND DIGESTIVE SYSTEM EDITED .docx",
    "CBT QUESTIONS.docx",
    "Paper II OBJECTIVES.docx",
    "2023 HOSPITAL FINAL PAPER II two.docx",
    "2023 HOSPITAL FINAL I  one (1).docx",
    "COUNCIL SAMPLE QUESTIONS.docx",
    "SET 45 MOCK EXAMS PAPER I OCTOBER 2023 (1).txt",
    "NMCN TUTORIALS (1).docx",
    "SET 45 MOCK EXAMS PAPER II OCTOBER 2023 (1).txt",
    "OSCE Questions1.docx",
    "OSCE Questions1 (1).docx",
    "Research and Bio MCQ.pdf",
    "200 pharmacology Original.pdf",
    "OSCE pass Q&A 2.pdf",
    "OSCE pass Q&A 2 (1).pdf",
    "Osce procedures mr sirja.pdf",
    "OSCE GUIDE.pdf",
    "NMCN objectives solved Pass questions.pdf",
    "RH mcq (1).pdf",
    "nmcn MAY 2024 NMCN CBT COMPLETEE.pdf",
    "Pharm Mcq.pdf",
    "nclex-rn-practice-questions-exam-cram-5th-edition.pdf",
    "Nursing-and-midwifery-council-past-questions-and-answers-for-Nursing-student (1).pdf",
    "SET 45 MOCK EXAMS PAPER II OCTOBER 2023.txt",
    "SET 45 MOCK EXAMS PAPER I OCTOBER 2023.txt",
    "2023 HOSPITAL FINAL I  one.docx",
    "Hospital Finals Paper 1-1.pdf",
    "p2 questions.docx",
    "250  Hospital final paper I QUESTIONS.docx",
    "NMCN TUTORIALS.docx",
]


TEXTBOOK_GUIDE = {
    "group": "Textbook exam strategy",
    "source": "Delmar's Practice Questions for NCLEX-PN, Chapter 1",
    "title": "How to pass CBT-style nursing council exams",
    "points": [
        "Read the whole stem first, then identify the client, problem, exact task, and time frame.",
        "Treat words such as first, initial, priority, best, most, least, early, and late as decision signals.",
        "For initial action, think assessment first unless airway, breathing, circulation, safety, or infection control is immediately threatened.",
        "Use Maslow and safety: keep the client breathing, keep the client safe, then address comfort and teaching.",
        "Eliminate options that do not answer the question, then choose the option most directly supported by the data in the stem.",
        "Practice many timed questions and review rationales for both correct and incorrect options.",
    ],
}


def file_hash(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def read_docx(path):
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    return "\n".join(paragraphs + table_lines)


def read_pdf(path):
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def read_text(path):
    if path.suffix.lower() == ".docx":
        return read_docx(path)
    if path.suffix.lower() == ".pdf":
        return read_pdf(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def normalize(text):
    text = text.replace("\u00a0", " ")
    text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def compact(text):
    return re.sub(r"\s+", " ", text).strip()


def parse_answer(raw):
    raw = raw.strip().upper()
    letters = re.findall(r"\b[A-H]\b", raw)
    if letters:
        return [ord(letter) - ord("A") for letter in letters]
    nums = re.findall(r"\b[1-8]\b", raw)
    return [int(num) - 1 for num in nums]


def source_category(source, prompt):
    blob = f"{source} {prompt}".lower()
    if any(term in blob for term in ["pharm", "drug", "medication", "insulin", "digoxin", "antibiotic", "dose"]):
        return "Pharmacological Therapies"
    if any(term in blob for term in ["nutrition", "diet", "digestive", "gastro", "liver", "bowel"]):
        return "Basic Care and Comfort"
    if any(term in blob for term in ["preg", "midwif", "reproductive", "antenatal", "postpartum", "labour", "labor", "newborn"]):
        return "Health Promotion and Maintenance"
    if any(term in blob for term in ["osce", "procedure", "sterile", "dressing", "catheter", "operation"]):
        return "Reduction of Risk Potential"
    if any(term in blob for term in ["research", "bio", "anatomy", "physiology"]):
        return "Physiological Adaptation"
    if any(term in blob for term in ["mental", "psychiatric", "depress", "anxiety"]):
        return "Psychosocial Integrity"
    if any(term in blob for term in ["infection", "safety", "isolation", "hand"]):
        return "Safety and Infection Prevention and Control"
    return "Coordinated Care" if "paper ii" in blob else "Physiological Adaptation"


def parse_question_chunks(text, source):
    lines = [line.strip() for line in normalize(text).splitlines() if line.strip()]
    chunks, current = [], []
    for line in lines:
        answer_match = re.match(r"^(?:ANS(?:WER)?|ANSWER)\s*[:=]?\s*([A-H1-8,\s]+)\.?$", line, re.I)
        if answer_match:
            if current:
                chunks.append((current, answer_match.group(1)))
                current = []
        else:
            current.append(line)

    return build_items_from_chunks(chunks, source)


def build_items_from_chunks(chunks, source):
    items = []
    for lines_before_answer, answer_raw in chunks:
        option_positions = []
        for idx, line in enumerate(lines_before_answer):
            if re.match(r"^(?:[_\s]*)(?:[A-Ha-h]|[1-8])[\).]\s+", line):
                option_positions.append(idx)

        if len(option_positions) >= 2:
            first_option = option_positions[0]
            prompt = compact(" ".join(lines_before_answer[:first_option]))
            options = []
            for pos, idx in enumerate(option_positions):
                end = option_positions[pos + 1] if pos + 1 < len(option_positions) else len(lines_before_answer)
                raw = " ".join(lines_before_answer[idx:end])
                raw = re.sub(r"^(?:[_\s]*)(?:[A-Ha-h]|[1-8])[\).]\s+", "", raw).strip()
                options.append(compact(raw))
        elif len(lines_before_answer) >= 5:
            prompt = compact(" ".join(lines_before_answer[:-4]))
            options = [compact(line) for line in lines_before_answer[-4:]]
        else:
            continue

        prompt = re.sub(r"^\d{1,4}[\).]\s*", "", prompt).strip()
        prompt = re.sub(r"^(?:QUESTION\s*)?\d{1,4}\s+", "", prompt, flags=re.I).strip()
        if len(prompt) < 12:
            continue

        answer = [index for index in parse_answer(answer_raw) if 0 <= index < len(options)]
        if not answer or len(options) < 2:
            continue

        items.append(
            {
                "source": source,
                "chapter": source,
                "category": source_category(source, prompt),
                "type": "multi" if len(answer) > 1 else "single",
                "prompt": prompt,
                "options": options[:8],
                "answer": answer,
                "rationale": f"Source answer: {', '.join(str(i + 1) for i in answer)}. This item was imported from {source}; no detailed rationale was supplied in the source document.",
            }
        )
    return items


def parse_marking_key_questions(text, source):
    normalized = normalize(text)
    split = re.split(r"\bMARKING\s+(?:GUIDE|SCHEME)\b", normalized, flags=re.I)
    if len(split) < 2:
        return []
    question_text, key_text = split[0], split[1]
    keys = re.findall(r"\b[A-H]\b", key_text.upper())
    if not keys:
        return []

    lines = [line.strip() for line in question_text.splitlines() if line.strip()]
    question_starts = [idx for idx, line in enumerate(lines) if re.match(r"^\d{1,4}[\).]\s+", line)]
    chunks = []
    if question_starts:
        for pos, start in enumerate(question_starts):
            end = question_starts[pos + 1] if pos + 1 < len(question_starts) else len(lines)
            if pos < len(keys):
                chunks.append((lines[start:end], keys[pos]))
    else:
        cursor = 0
        for answer in keys:
            if cursor + 5 > len(lines):
                break
            chunks.append((lines[cursor : cursor + 5], answer))
            cursor += 5

    return build_items_from_chunks(chunks, source)


def guide_group_for(name):
    lower = name.lower()
    if "osce" in lower:
        return "OSCE guide"
    if "tutorial" in lower or "guide" in lower or "objectives" in lower:
        return "Council/NMCN guide"
    if "nutrition" in lower or "pharm" in lower or "research" in lower or "bio" in lower:
        return "Subject guide"
    return "Past-question guide"


def extract_guide_cards(name, text):
    lower = name.lower()
    guideish = any(key in lower for key in ["tutorial", "guide", "objectives", "osce", "nutrition", "pharm", "research", "bio"])
    if not guideish:
        return []

    paragraphs = [compact(p) for p in re.split(r"\n\s*\n|\r\n\s*\r\n", normalize(text)) if len(compact(p)) > 45]
    clean = []
    for paragraph in paragraphs:
        if re.search(r"^(?:\d+[\).]|[A-D][\).]|ANSWER\s*:)", paragraph, re.I):
            continue
        if "CamScanner" in paragraph:
            continue
        clean.append(paragraph)

    cards = []
    for paragraph in clean[:12]:
        sentences = re.split(r"(?<=[.!?])\s+", paragraph)
        points = [s.strip() for s in sentences if len(s.strip()) > 20][:5]
        if points:
            cards.append(
                {
                    "group": guide_group_for(name),
                    "source": name,
                    "title": points[0][:90],
                    "points": points,
                }
            )
    return cards


def item_key(item):
    return compact(item["prompt"] + " " + " ".join(item["options"])).lower()


def main():
    seen_files = set()
    seen_items = set()
    questions = []
    guides = [TEXTBOOK_GUIDE]
    file_report = []

    for name in FILES:
        path = DOWNLOADS / name
        if not path.exists():
            continue
        digest = file_hash(path)
        if digest in seen_files:
            file_report.append({"file": name, "status": "duplicate skipped"})
            continue
        seen_files.add(digest)

        text = read_text(path)
        parsed = parse_question_chunks(text, name)
        parsed.extend(parse_marking_key_questions(text, name))
        added = 0
        for item in parsed:
            key = item_key(item)
            if key in seen_items:
                continue
            seen_items.add(key)
            item["id"] = f"supp-{len(questions) + 1}"
            questions.append(item)
            added += 1

        guide_cards = extract_guide_cards(name, text)
        guides.extend(guide_cards)
        file_report.append({"file": name, "questions": added, "guides": len(guide_cards), "chars": len(text)})

    QUESTION_OUT.write_text(
        "/* Generated from user-supplied council past-question documents. */\n"
        f"export const supplementalQuestions = {json.dumps(questions, ensure_ascii=False, indent=2)};\n",
        encoding="utf-8",
    )
    GUIDE_OUT.write_text(
        "/* Generated guide cards from the textbook and user-supplied council materials. */\n"
        f"export const guideSections = {json.dumps(guides, ensure_ascii=False, indent=2)};\n",
        encoding="utf-8",
    )
    print(json.dumps({"questions": len(questions), "guide_cards": len(guides), "files": file_report}, indent=2))


if __name__ == "__main__":
    main()
